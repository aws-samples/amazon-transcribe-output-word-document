"""
This sample, non-production-ready application will produce Word Document transcriptions using automatic speech
recognition from Amazon Bedrock Data Automation, with or with the standard customer blueprint.  The a pplication
requires the following non-standard python libraries to be installed:

- python-docx
- scipy
- matplotlib

© 2026 Amazon Web Services, Inc. or its affiliates. All Rights Reserved.
This AWS Content is provided subject to the terms of the AWS Customer Agreement available at
http://aws.amazon.com/agreement or other written agreement between Customer and either
Amazon Web Services, Inc. or Amazon Web Services EMEA SARL or both.
"""

from docx import Document
from docx.shared import Cm, Mm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from pathlib import Path
from time import perf_counter
from scipy.interpolate import make_interp_spline
import urllib.request
import json
import datetime
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import numpy as np
import statistics
import os
import boto3
import argparse
from io import BytesIO


# Common formats and styles
CUSTOM_STYLE_HEADER = "CustomHeader"
TABLE_STYLE_STANDARD = "Light List"
CATEGORY_TRANSCRIPT_BG_COLOUR = "EEFFFF"
CATEGORY_TRANSCRIPT_FG_COLOUR = RGBColor(0, 128, 255)
SMALL_TEXT = Pt(7)
ALTERNATE_ROW_COLOUR = "F0F0F0"
BAR_CHART_WIDTH = 1.0

# Column offsets in Transcribe output document table
COL_STARTTIME = 0
COL_DURATION = 1
COL_SPEAKER = 2
COL_SENTIMENT = 3
COL_CONTENT = 4

# Comprehend Sentiment helpers - note, if a language code in Comprehend has multiple suffixed versions
# then the suffixed versions MUST be defined in the language list BEFORE the base one; e.h. "zh-TW" before "zh"
MIN_SENTIMENT_LENGTH = 16
MIN_SENTIMENT_NEGATIVE = 0.4
MIN_SENTIMENT_POSITIVE = 0.6
SENTIMENT_LANGUAGES = ["en", "es", "fr", "de", "it", "pt", "ar", "hi", "ja", "ko", "zh-TW", "zh"]

# Image download URLS
IMAGE_URL_BANNER = "https://raw.githubusercontent.com/aws-samples/amazon-transcribe-output-word-document/main/images/banner-bda.png"
IMAGE_URL_SMILE = "https://raw.githubusercontent.com/aws-samples/amazon-transcribe-output-word-document/main/images/smile.png"
IMAGE_URL_FROWN = "https://raw.githubusercontent.com/aws-samples/amazon-transcribe-output-word-document/main/images/frown.png"
IMAGE_URL_NEUTRAL = "https://raw.githubusercontent.com/aws-samples/amazon-transcribe-output-word-document/main/images/neutral.png"


# Additional Constants
START_NEW_SEGMENT_DELAY = 2.0       # After n seconds pause by one speaker, put next speech in new segment

class SpeechSegment:
    """ Class to hold information about a single speech segment """
    def __init__(self):
        self.segmentStartTime = 0.0
        self.segmentEndTime = 0.0
        self.segmentSpeaker = ""
        self.segmentText = ""
        self.segmentLanguage = ""
        self.segmentConfidence = []
        self.segmentContentModeration = []
        self.segmentSentimentScore = -1.0    # -1.0 => no sentiment calculated
        self.segmentPositive = 0.0
        self.segmentNegative = 0.0
        self.segmentIsPositive = False
        self.segmentIsNegative = False

def convert_timestamp(time_in_seconds):
    """
    Function to help convert timestamps from s to H:M:S:MM

    :param time_in_seconds: Time in seconds to be displayed
    :return: Formatted string for this timestamp value
    """
    timeDelta = datetime.timedelta(seconds=float(time_in_seconds))
    tsFront = timeDelta - datetime.timedelta(microseconds=timeDelta.microseconds)
    tsSmall = timeDelta.microseconds
    return str(tsFront) + "." + str(int(tsSmall / 10000))


def set_transcript_text_style(run, force_highlight, confidence=0.0, rgb_color=None):
    """
    Sets the colour and potentially the style of a given run of text in a transcript.  You can either
    supply the hex-code, or base it upon the confidence score in the transcript.

    :param run: DOCX paragraph run to be modified
    :param force_highlight: Indicates that we're going to forcibly set the background colour
    :param confidence: Confidence score for this word, used to dynamically set the colour
    :param rgb_color: Specific colour for the text
    """

    # If we have an RGB colour then use it
    if rgb_color is not None:
        run.font.color.rgb = rgb_color
    else:
        # Set the colour based upon the supplied confidence score
        if confidence >= 0.90:
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif confidence >= 0.5:
            run.font.color.rgb = RGBColor(102, 51, 0)
        else:
            run.font.color.rgb = RGBColor(255, 0, 0)

    # Apply any other styles wanted
    if confidence == 0.0:
        # Call out any total disasters in bold
        run.font.bold = True

    # Force the background colour if required
    if force_highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def write_transcribe_text(output_table, cli_args, speech_segments, timed_topics):
    """
    Writes out each line of the transcript in the Word table structure, optionally including sentiments

    :param output_table: Word document structure to write the table into
    :param cli_args: Our entrypoint CLI arguments
    :param speech_segments: Turn-by-turn speech list
    :param timed_topics: List of categories identified at any timestamps
    """
    topics_unmarked = []

    # Load our image files if we have sentiment enabled
    if cli_args.sentiment == 'on':
        png_smile = load_image(IMAGE_URL_SMILE)
        png_frown = load_image(IMAGE_URL_FROWN)
        png_neutral = load_image(IMAGE_URL_NEUTRAL)
        content_col_offset = 0
    else:
        # Ensure we offset the CONTENT column correctly due to no sentiment
        content_col_offset = -1

    # Get our topic timers, if there are any
    if timed_topics != []:
        topics_unmarked = [d["start_time"] for d in timed_topics]

    # Create a row populate it for each segment that we have
    shading_reqd = False
    for segment in speech_segments:
        # Start with the easy stuff
        start_in_millis = int(segment.segmentStartTime * 1000.0)
        end_in_millis = int(segment.segmentEndTime * 1000.0)
        row_cells = output_table.add_row().cells
        row_cells[COL_STARTTIME].text = convert_timestamp(segment.segmentStartTime)
        row_cells[COL_DURATION].text = f"{(segment.segmentEndTime - segment.segmentStartTime):.1f}s"
        row_cells[COL_SPEAKER].text = segment.segmentSpeaker

        # Start this line with a topic marker (or more) if we need to
        while start_in_millis in topics_unmarked:
            # Mark this one as being the start of a topic
            idx = topics_unmarked.index(start_in_millis)
            text_to_add = "[Topic "
            text_to_add += str(timed_topics[idx]["index"] + 1)
            text_to_add += "]"
            run = row_cells[COL_CONTENT + content_col_offset].paragraphs[0].add_run(text_to_add)
            run.font.size = SMALL_TEXT
            set_transcript_text_style(run, False, confidence=0.0)
            row_cells[COL_CONTENT + content_col_offset].paragraphs[0].add_run(" ")

            # Remove from the lists in case there's 2+ starting at this point
            timed_topics.pop(idx)
            topics_unmarked.pop(idx)

        # Then do each word with confidence-level colouring
        text_index = 1
        for eachWord in segment.segmentConfidence:
            # Output the next word, with the correct confidence styling and forced background
            run = row_cells[COL_CONTENT + content_col_offset].paragraphs[0].add_run(eachWord["text"])
            text_index += len(eachWord["text"])
            confLevel = eachWord["confidence"]
            set_transcript_text_style(run, False, confidence=confLevel)

        # Check if we have to add our guardrail results
        if cli_args.guardrailCheck == 'on':
               if segment.segmentContentModeration != []:
                    # There is something here - start with the header
                    run = row_cells[COL_CONTENT + content_col_offset].paragraphs[0].add_run("\n")
                    run = row_cells[COL_CONTENT + content_col_offset].paragraphs[0].add_run(['Guardrail'])
                    set_transcript_text_style(run, True, confidence=0.0)
                    run.font.size = SMALL_TEXT
                    run.font.italic = True
                    run = row_cells[COL_CONTENT + content_col_offset].paragraphs[0].add_run(" ")
                    run.font.size = SMALL_TEXT
                    run.font.italic = True

                    # Loop round each logged guardrail trigged against this segment
                    for guardrailTrigger in segment.segmentContentModeration:
                        text_to_add = guardrailTrigger['category'] + " (" + str(int(guardrailTrigger['confidence'] * 100.0)) + "%)  " 
                        run = row_cells[COL_CONTENT + content_col_offset].paragraphs[0].add_run(text_to_add)
                        run.font.size = SMALL_TEXT
                        run.font.italic = True

        # If enabled, finish with the base sentiment for the segment - don't write out
        # score if it turns out that this segment ie neither Negative nor Positive
        if cli_args.sentiment == 'on':
            if segment.segmentIsPositive or segment.segmentIsNegative:
                paragraph = row_cells[COL_SENTIMENT].paragraphs[0]
                img_run = paragraph.add_run()
                if segment.segmentIsPositive:
                    img_run.add_picture(png_smile, width=Mm(4))
                else:
                    img_run.add_picture(png_frown, width=Mm(4))

                # We only have turn-by-turn sentiment score values in non-analytics mode
                # text_run = paragraph.add_run(' (' + str(segment.segmentSentimentScore)[:4] + ')')
                # text_run.font.size = SMALL_TEXT
                # text_run.font.italic = True
            # else:
            #     row_cells[COL_SENTIMENT].paragraphs[0].add_run().add_picture(png_neutral, width=Mm(4))

        # Add highlighting to the row if required
        if shading_reqd:
            for column in range(0, COL_CONTENT + content_col_offset + 1):
                set_table_cell_background_colour(row_cells[column], ALTERNATE_ROW_COLOUR)
        shading_reqd = not shading_reqd

        # Before we end, does an analytics category start with this line's end time?
        if end_in_millis in timed_topics:
            # If so, write out the line after this
            break


def merge_speaker_segments(input_segment_list):
    """
    Merges together consecutive speaker segments unless:
    (a) There is a speaker change, or
    (b) The gap between segments is greater than our acceptable level of delay
    (c) The language changes

    :param input_segment_list: Full time-sorted list of speaker segments
    :return: An updated segment list
    """
    outputSegmentList = []
    lastSpeaker = ""
    lastLanguage = ""
    lastSegment = None

    # Step through each of our defined speaker segments
    for segment in input_segment_list:
        if (segment.segmentSpeaker != lastSpeaker) or \
                ((segment.segmentStartTime - lastSegment.segmentEndTime) >= START_NEW_SEGMENT_DELAY) or \
                (segment.segmentLanguage != lastLanguage) or \
                (segment.segmentContentModeration != []):
            # Simple case - speaker change or > n-second gap means new output segment
            outputSegmentList.append(segment)

            # This is now our base segment moving forward
            lastSpeaker = segment.segmentSpeaker
            lastLanguage = segment.segmentLanguage
            lastSegment = segment
        else:
            # Same speaker, short time, need to copy this info to the last one
            lastSegment.segmentEndTime = segment.segmentEndTime
            lastSegment.segmentText += " " + segment.segmentText
            segment.segmentConfidence[0]["text"] = " " + segment.segmentConfidence[0]["text"]
            for wordConfidence in segment.segmentConfidence:
                lastSegment.segmentConfidence.append(wordConfidence)

    return outputSegmentList


def generate_sentiment(segment_list):
    """
    Generates sentiment per speech segment, inserting the results into the input list.  This will use
    Amazon Comprehend, but we need to map the language code per line to one that Comprehend understands

    :param segment_list: List of speech segments
    """

    # Get our botot3 client, then go through each segment
    client = boto3.client("comprehend")
    for nextSegment in segment_list:
        # Only continue if Comprehent knows this language
        if nextSegment.segmentLanguage in SENTIMENT_LANGUAGES:
            # Only continue if it's a long enough text fragment
            if len(nextSegment.segmentText) >= MIN_SENTIMENT_LENGTH:
                nextText = nextSegment.segmentText
                response = client.detect_sentiment(Text=nextText, LanguageCode=nextSegment.segmentLanguage)
                positiveBase = response["SentimentScore"]["Positive"]
                negativeBase = response["SentimentScore"]["Negative"]

                # If we're over the NEGATIVE threshold then we're negative
                if negativeBase >= MIN_SENTIMENT_NEGATIVE:
                    nextSegment.segmentIsNegative = True
                    nextSegment.segmentSentimentScore = negativeBase
                # Else if we're over the POSITIVE threshold then we're positive,
                # otherwise we're either MIXED or NEUTRAL and we don't really care
                elif positiveBase >= MIN_SENTIMENT_POSITIVE:
                    nextSegment.segmentIsPositive = True
                    nextSegment.segmentSentimentScore = positiveBase

                # Store all of the original sentiments for future use
                nextSegment.segmentAllSentiments = response["SentimentScore"]
                nextSegment.segmentPositive = positiveBase
                nextSegment.segmentNegative = negativeBase


def set_repeat_table_header(row):
    """
    Set Word repeat table row on every new page.  Cannot be directorly done in python-docx
    """
    tbl_header = OxmlElement('w:tblHeader') # create new oxml element flag which indicates that row is header row
    first_row_props = row._element.get_or_add_trPr() # get if exists or create new table row properties el
    first_row_props.append(tbl_header) # now first row is the header row


def load_image(url):
    """
    Loads binary image data from a URL for later embedding into a docx document
    :param url: URL of image to be downloaded
    :return: BytesIO object that can be added as a docx image
    """
    image_url = urllib.request.urlopen(url)
    io_url = BytesIO()
    io_url.write(image_url.read())
    io_url.seek(0)
    return io_url


def create_text_array(whole_array):
    """
    Iterates through an array and returns a single-line text representation of it

    :param whole_array: Array to iterate through
    :return: Single text value representing the whole array
    """
    index = 0
    text_line = ""

    for array_element in whole_array:
        # If we're the 2nd+ line then add some spaceing
        if index > 0:
            text_line += ", "
        text_line += array_element
        index += 1

    return text_line


def yes_or_no_icon(check):
    """
    Writes a consistent icon for a boolean value
    :param check: Boolean to indicate the check needed
    :return: Consistent checkmark icon
    """
    if check:
        return "✅"
    else:
        return "❌"


def create_pie_chart(document, temp_files, output_file_name, pie_value):
    """
    Creates a segmented pie chart with two segments sized 1..5, and both segments fill
    up the pie up to 5 chunks in size

    :param document: Document we're writing into
    :param temp_files: List of temporary files we're creating
    :param output_file_name: Name for this temporary file
    :param pie_value: Size (1..5) of the leading pie number
    """
    fig, ax = plt.subplots()

    size = 0.8
    sizes = [pie_value, 5-pie_value]

    second_colour_offset = [19, 4, 7, 3, 10, 8][int(pie_value)]
    tab20c = plt.color_sequences["tab20c"]
    outer_colors = [tab20c[i] for i in [second_colour_offset, 19]]

    ax.pie(sizes, radius=1.5, colors=outer_colors, startangle=90, wedgeprops=dict(width=size, edgecolor='w'))
    ax.text(0, 0, str(pie_value), ha='center', va='center', fontsize=36, fontweight='bold')

    chart_file_name = "./" + output_file_name + ".png"
    plt.savefig(chart_file_name)
    temp_files.append(chart_file_name)
    plt.clf()
    document.add_picture(chart_file_name, width=Cm(3.5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_section_columns(document, columns):
    """
    :param document: The document holding the final (our) section
    :param columns: Number of columns to set on the section
    """
    section_ptr = document.sections[-1]._sectPr
    cols = section_ptr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(columns))


def write(cli_arguments, speech_segments, custom_json):
    """
    Write a transcript from the .json transcription file and other data generated
    by the results parser, putting it all into a human-readable Word document

    :param cli_arguments: CLI arguments used for this processing run
    :param speech_segments: List of call speech segments
    :param custom_json: Custom JSON data associated with our original input file
    """

    json_filepath = Path(cli_arguments.inputFile)
    data = json.load(open(json_filepath.absolute(), "r", encoding="utf-8"))
    sentimentEnabled = (cli_arguments.sentiment == 'on')
    temp_files = []
    timed_topics = []

    # Initiate Document, orientation and margins
    document = Document()
    document.sections[0].left_margin = Mm(19.1)
    document.sections[0].right_margin = Mm(19.1)
    document.sections[0].top_margin = Mm(19.1)
    document.sections[0].bottom_margin = Mm(19.1)
    document.sections[0].page_width = Mm(210)
    document.sections[0].page_height = Mm(297)

    # Set the base font and document title
    font = document.styles["Normal"].font
    font.name = "Calibri"
    font.size = Pt(9)

    # Create our custom text header style
    custom_style = document.styles.add_style(CUSTOM_STYLE_HEADER, WD_STYLE_TYPE.PARAGRAPH)
    custom_style.paragraph_format.widow_control = True
    custom_style.paragraph_format.keep_with_next = True
    custom_style.paragraph_format.space_after = Pt(0)
    custom_style.font.size = font.size
    custom_style.font.name = font.name
    custom_style.font.bold = True
    custom_style.font.italic = True

    # Intro banner header
    document.add_picture(load_image(IMAGE_URL_BANNER), width=Mm(171))

    # We need 2 columns only if we're in custom mode, as we put the custom summary on the right of the table
    if custom_json != {}:
        document.add_section(WD_SECTION.CONTINUOUS)
        set_section_columns(document, 2)

    # Write out the call summary table
    table = document.add_table(rows=1, cols=2)
    table.style = document.styles[TABLE_STYLE_STANDARD]
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Amazon BDA Audio Metadata"
    hdr_cells[0].merge(hdr_cells[1])

    # Now add in the data row by row
    job_data = []
    job_data.append({"name": "Audio Filename", "value": data["metadata"]["s3_key"]})
    job_data.append({"name": "Audio Duration", "value": convert_timestamp(data["metadata"]["duration_millis"] / 1000)})
    output_data = str(data["metadata"]["sample_rate"])
    output_data += "kHz "
    output_data +=  data["metadata"]["format"] + "-file"
    job_data.append({"name": "Audio Format", "value": str(output_data)})
    job_data.append({"name": "Dominant Language", "value": data["metadata"]["dominant_asset_language"]})
    job_data.append({"name": "Generative Language", "value": data["metadata"]["generative_output_language"]})
   
    # Place all of our job-summary fields into the Table, one row at a time
    for next_row in job_data:
        row_cells = table.add_row().cells
        row_cells[0].text = next_row["name"]
        row_cells[1].text = next_row["value"]

    # Formatting transcript table widths
    widths = (Cm(3.44), Cm(4.89))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Spacer paragraph
    document.add_paragraph()

    # Put any custom fields on the RHS of the page
    if custom_json != {}:
        # Write out the custom metadata table
        document.add_section(WD_SECTION.NEW_COLUMN)
        table = document.add_table(rows=1, cols=2)
        table.style = document.styles[TABLE_STYLE_STANDARD]
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Amazon BDA Audio Custom Metadata"
        hdr_cells[0].merge(hdr_cells[1])

        # Add in each repeatable block
        job_data = []
        job_data.append({"name": "Categories", "value": create_text_array(custom_json["inference_result"]["call_categories"])})
        job_data.append({"name": "Topics", "value": create_text_array(custom_json["inference_result"]["call_topics"])})
        job_data.append({"name": "Issues", "value": create_text_array(custom_json["inference_result"]["call_issues"])})
        job_data.append({"name": "Intents", "value": create_text_array(custom_json["inference_result"]["call_intents"])})
        job_data.append({"name": "Agent Actions", "value": create_text_array(custom_json["inference_result"]["agent_actions"])})
        job_data.append({"name": "Pending Actions", "value": create_text_array(custom_json["inference_result"]["agent_pending_action_items"])})

        # Place all of our job-summary fields into the Table, one row at a time
        for next_row in job_data:
            row_cells = table.add_row().cells
            row_cells[0].text = next_row["name"]
            row_cells[1].text = next_row["value"]

        # Formatting transcript table widths
        widths = (Cm(3.0), Cm(5.1))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

    # New single-column section
    document.add_section(WD_SECTION.CONTINUOUS)
    set_section_columns(document, 1)

    # At this point, if we have no transcript then we need to quickly exit
    if len(speech_segments) == 0:
        write_custom_text_header(document, "This call had no audible speech to transcribe.")
    else:
        # Write out any call summarisation data
        if (custom_json != {}) and ("call_summary" in custom_json["inference_result"]):
            # The custom version is shorter, so let's use it if we can
            write_detected_summaries(document, custom_json["inference_result"]["call_summary"])
        elif "summary" in data["audio"]:
            # The standard call summary in BDA may not exist if the call isn't summarisable
            write_detected_summaries(document, data["audio"]["summary"])

        # Write out our custom table on how the call went, along with our 1-5 charts
        if custom_json != {}:
            # Start with our 1-5 bars
            document.add_section(WD_SECTION.CONTINUOUS)
            set_section_columns(document, 3)

            # Do all the bars
            write_custom_text_header(document, 'Customer Satisfaction', position=WD_ALIGN_PARAGRAPH.CENTER)
            create_pie_chart(document, temp_files, "yes_no_chart_1", int(custom_json["inference_result"]["caller_satisfaction_level"]))
            document.add_section(WD_SECTION.NEW_COLUMN)
            write_custom_text_header(document, 'Caller Emotion', position=WD_ALIGN_PARAGRAPH.CENTER)
            create_pie_chart(document, temp_files, "yes_no_chart_2", int(custom_json["inference_result"]["caller_emotion_rating"]))
            document.add_section(WD_SECTION.NEW_COLUMN)
            write_custom_text_header(document, 'Agent Emotion', position=WD_ALIGN_PARAGRAPH.CENTER)
            create_pie_chart(document, temp_files, "yes_no_chart_3", int(custom_json["inference_result"]["agent_emotion_rating"]))

            # Output other tables - start with how the call went
            document.add_section(WD_SECTION.CONTINUOUS)
            set_section_columns(document, 3)

            table = document.add_table(rows=1, cols=2, style=TABLE_STYLE_STANDARD)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Call Success"
            hdr_cells[0].merge(hdr_cells[1])

            # Add in each repeatable block
            job_data = []
            job_data.append({"name": "Call Issue Resolved", "value": yes_or_no_icon(custom_json["inference_result"]["issue_resolution"])})
            job_data.append({"name": "Call Opening Used", "value": yes_or_no_icon(custom_json["inference_result"]["call_opening"])})
            job_data.append({"name": "Call Wrapup Used", "value": yes_or_no_icon(custom_json["inference_result"]["call_wrapup"])})
            job_data.append({"name": "Caller Neg Emotion", "value": yes_or_no_icon(custom_json["inference_result"]["caller_negative_emotion"])})

            # Place all of our job-summary fields into the Table, one row at a time
            for next_row in job_data:
                row_cells = table.add_row().cells
                row_cells[0].text = next_row["name"]
                row_cells[1].text = next_row["value"]

            # Formatting transcript table widths
            widths = [Inches(3.0), Inches(1.0)]
            for row in table.rows:
                row.cells[0].width = widths[0]
                row.cells[1].width = widths[1]

            # Finish off spacing
            document.add_paragraph()

            # Next block is the caller sentiment table
            document.add_section(WD_SECTION.NEW_COLUMN)

            table = document.add_table(rows=1, cols=2, style=TABLE_STYLE_STANDARD)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Caller Sentiment"
            hdr_cells[0].merge(hdr_cells[1])
            row_cells = table.add_row().cells
            row_cells[0].text = custom_json["inference_result"]["caller_sentiment_summary"]
            row_cells[0].merge(row_cells[1])
            row_cells[0].paragraphs[0].runs[0].font.bold = False

            # Add in each repeatable block
            job_data = []
            job_data.append({"name": "Emotion Label", "value": custom_json["inference_result"]["caller_emotion_label"]})
            job_data.append({"name": "End Sentiment", "value": custom_json["inference_result"]["caller_end_sentiment"]})
            job_data.append({"name": "Improvement", "value": custom_json["inference_result"]["caller_emotion_improvement"]})

            # Place all of our job-summary fields into the Table, one row at a time
            for next_row in job_data:
                row_cells = table.add_row().cells
                row_cells[0].text = next_row["name"]
                row_cells[1].text = next_row["value"]

            # Formatting transcript table widths
            widths = [Inches(2.5), Inches(1.5)]
            for row in table.rows:
                row.cells[0].width = widths[0]
                row.cells[1].width = widths[1]

            # Next block is the agent sentiment table
            document.add_section(WD_SECTION.NEW_COLUMN)

            table = document.add_table(rows=1, cols=2, style=TABLE_STYLE_STANDARD)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Agent Sentiment"
            hdr_cells[0].merge(hdr_cells[1])
            row_cells = table.add_row().cells
            row_cells[0].text = custom_json["inference_result"]["agent_sentiment_summary"]
            row_cells[0].merge(row_cells[1])
            row_cells[0].paragraphs[0].runs[0].font.bold = False

            # Add in each repeatable block
            job_data = []
            job_data.append({"name": "Emotion Label", "value": custom_json["inference_result"]["agent_emotion_label"]})
            job_data.append({"name": "End Sentiment", "value": custom_json["inference_result"]["agent_end_sentiment"]})

            # Place all of our job-summary fields into the Table, one row at a time
            for next_row in job_data:
                row_cells = table.add_row().cells
                row_cells[0].text = next_row["name"]
                row_cells[1].text = next_row["value"]

            # Formatting transcript table widths
            widths = [Inches(2.5), Inches(1.5)]
            for row in table.rows:
                row.cells[0].width = widths[0]
                row.cells[1].width = widths[1]

            document.add_section(WD_SECTION.CONTINUOUS)

        # Work out where our topics happened
        if "topics" in data:
            for topic in data["topics"]:
                timed_topics.append({"start_time": topic["start_timestamp_millis"], "index": topic["topic_index"]})

        # Process and display transcript by speaker segments (new section)
        # -- Conversation "turn" start time and duration
        # -- Speaker identification
        # -- Sentiment type (if enabled) and sentiment score (if available)
        # -- Transcribed text with (if available) Call Analytics markers
        document.add_section(WD_SECTION.CONTINUOUS)
        set_section_columns(document, 1)
        write_custom_text_header(document, "Call Transcription")
        table_cols = 4
        if sentimentEnabled:
            # Ensure that we add space for the sentiment column
            table_cols += 1
            content_col_offset = 0
        else:
            # Will need to shift the content column to the left, as Sentiment isn't there now
            content_col_offset = -1
        table = document.add_table(rows=1, cols=table_cols)
        table.style = document.styles[TABLE_STYLE_STANDARD]
        hdr_cells = table.rows[0].cells
        hdr_cells[COL_STARTTIME].text = "Start"
        hdr_cells[COL_DURATION].text = "Dur."
        hdr_cells[COL_SPEAKER].text = "Speaker"
        hdr_cells[COL_CONTENT + content_col_offset].text = "Transcription"

        # Based upon our segment list, write out the transcription table
        write_transcribe_text(table, cli_arguments, speech_segments, timed_topics)
        document.add_paragraph()

        # Formatting transcript table widths - we need to add sentiment
        # column if needed, and it and the content width accordingly
        widths = [Inches(0.8), Inches(0.5), Inches(0.5), 0]
        if sentimentEnabled:
            # Comprehend sentiment needs space for the icon and % score
            widths.append(0)
            widths[COL_CONTENT + + content_col_offset] = Inches(7.3)
            widths[COL_SENTIMENT] = Inches(0.3)
        else:
            widths[COL_CONTENT + content_col_offset] = Inches(7.6)

        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        # Setup the repeating header
        set_repeat_table_header(table.rows[0])

        # Write out an call topics data
        if "topics" in data:
            write_topic_reasons(document, data["topics"])

        # Generate our raw data for the Comprehend sentiment graph (if requested)
        if sentimentEnabled:
            write_comprehend_sentiment(document, speech_segments, temp_files)

    # Save the whole document
    document.save(cli_arguments.outputFile)

    # Now delete any local images that we created
    for filename in temp_files:
        os.remove(filename)


def write_custom_text_header(document, text_label, position=WD_ALIGN_PARAGRAPH.LEFT):
    """
    Adds a run of text to the document with the given text label, but using our customer text-header style

    :param document: Word document structure to write the table into
    :param text_label: Header text to write out
    :paran position: Optional parameter to speficy the paragraph align method
    :return:
    """
    paragraph = document.add_paragraph(text_label)
    paragraph.style = CUSTOM_STYLE_HEADER
    paragraph.paragraph_format.alignment = position


def write_topic_reasons(document, topics):
    """
    Outputs the detected topic summaries

    :param document: Word document structure to write the table into
    :param topics: Topics structures
    """

    # Start with a two-columm table
    document.add_section(WD_SECTION.CONTINUOUS)
    table = document.add_table(rows=1, cols=2, style=TABLE_STYLE_STANDARD)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Conversational Topics"
    hdr_cells[0].merge(hdr_cells[1])

    # Output the generated topic summary reasons
    for topic in topics:
        row_cells = table.add_row().cells
        row_cells[0].text = str(topic["topic_index"]+1)
        row_cells[1].text = topic["summary"]

    # Formatting transcript table widths
    widths = [Inches(0.3), Inches(9.2)]
    for row in table.rows:
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]

    # Setup the repeating header
    set_repeat_table_header(table.rows[0])

    # Finish off spacing
    document.add_paragraph()


def write_detected_summaries(document, summary_text):
    """
    Outputs the detected speed summary on its own

    :param document: Word document structure to write the table into
    :param summary_text: Call summary text
    """

    # Start with a new single-column section
    document.add_section(WD_SECTION.CONTINUOUS)
    table = document.add_table(rows=1, cols=1, style="Table Grid")
    hdr_cells = table.rows[0].cells
    set_table_cell_background_colour(hdr_cells[0], '000000')
    parapraph = hdr_cells[0].paragraphs[0]
    run = parapraph.add_run("Call Summary Highlights")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    row_cells = table.add_row().cells
    row_cells[0].text = summary_text

    # Finish off with some spacing
    document.add_paragraph()

def write_comprehend_sentiment(document, speech_segments, temp_files):
    """
    Writes out a compound graph on per-speaker sentiment

    :param document: Docx document to add the sentiment graph to
    :param speech_segments: Process transcript text holding turn-by-turn sentiment
    :param temp_files: List of temp files to be deleted later
    :return:
    """
    # Initialise our base structures
    speaker0labels = ['ch_0', 'spk_0']
    speaker1labels = ['ch_1', 'spk_1']
    speaker0timestamps = []
    speaker0data = []
    speaker1timestamps = []
    speaker1data = []

    # Start with some spacing and a new sub-header
    document.add_paragraph()
    write_custom_text_header(document, "Amazon Comprehend Sentiment")
    # Now step through and process each speech segment's sentiment
    for segment in speech_segments:
        if segment.segmentIsPositive or segment.segmentIsNegative:
            # Only interested in actual sentiment entries
            if segment.segmentIsNegative:
                score = -segment.segmentSentimentScore
            else:
                score = segment.segmentSentimentScore
            timestamp = segment.segmentStartTime

            if segment.segmentSpeaker in speaker1labels:
                speaker1data.append(score)
                speaker1timestamps.append(timestamp)
            elif segment.segmentSpeaker in speaker0labels:
                speaker0data.append(score)
                speaker0timestamps.append(timestamp)

    # Spline fit needs at least 4 points for k=3, but 5 works better
    speaker1k = 3
    speaker0k = 3
    if len(speaker1data) < 5:
        speaker1k = 1
    if len(speaker0data) < 5:
        speaker0k = 1

    # Create Speaker-0 graph
    plt.figure(figsize=(8, 5))
    speaker0xnew = np.linspace(speaker0timestamps[0], speaker0timestamps[-1],
                               int((speaker0timestamps[-1] - speaker0timestamps[0]) + 1.0))
    speaker0spl = make_interp_spline(speaker0timestamps, speaker0data, k=speaker0k)
    speaker0powerSmooth = speaker0spl(speaker0xnew)
    plt.plot(speaker0timestamps, speaker0data, "ro")
    plt.plot(speaker0xnew, speaker0powerSmooth, "r", label="Speaker 0")

    # Create Speaker-1 graph
    speaker1xnew = np.linspace(speaker1timestamps[0], speaker1timestamps[-1],
                               int((speaker1timestamps[-1] - speaker1timestamps[0]) + 1.0))
    speaker1spl = make_interp_spline(speaker1timestamps, speaker1data, k=speaker1k)
    speaker1powerSmooth = speaker1spl(speaker1xnew)
    plt.plot(speaker1timestamps, speaker1data, "bo")
    plt.plot(speaker1xnew, speaker1powerSmooth, "b", label="Speaker 1")

    # Draw it out
    plt.title("Call Sentiment - Pos/Neg Only")
    plt.xlabel("Time (seconds)")
    plt.axis([0, max(speaker0timestamps[-1], speaker1timestamps[-1]), -1.1, 1.25])
    plt.legend()
    plt.axhline(y=0, color='k')
    plt.axvline(x=0, color='k')
    plt.grid(True)
    plt.xticks(np.arange(0, max(speaker0timestamps[-1], speaker1timestamps[-1]), 60))
    plt.yticks(np.arange(-1, 1.01, 0.25))

    # Write out the chart
    chart_file_name = "./" + "sentiment.png"
    plt.savefig(chart_file_name)
    temp_files.append(chart_file_name)
    plt.clf()
    document.add_picture(chart_file_name, width=Cm(14.64))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_table_cell_background_colour(cell, rgb_hex):
    """
    Modifies the background color of the given table cell to the given RGB hex value.  This currently isn't
    supporting by the DOCX module, and the only option is to modify the underlying Word document XML

    :param cell: Table cell to be changed
    :param rgb_hex: RBG hex string for the background color
    """
    parsed_xml = parse_xml(r'<w:shd {0} w:fill="{1}"/>'.format(nsdecls('w'), rgb_hex))
    cell._tc.get_or_add_tcPr().append(parsed_xml)


def create_turn_by_turn_segments(data, cli_args):
    """
    This creates a list of per-turn speech segments based upon the BDA data

    :param data: JSON result data from Transcribe
    :param cli_args: CLI arguments used for this processing run
    :return: List of transcription speech segments
    :return: Flag to indicate the presence of call summary data
    """
    speechSegmentList = []
    lastSpeaker = ""
    lastEndTime = 0.0
    skipLeadingSpace = False
    confidenceList = []
    nextSpeechSegment = None
    index = 0

    # Each turn has already been processed by BDA, so the outputs are in order
    for turn in data["audio_segments"]:

        # Setup the next speaker block
        nextSpeechSegment = SpeechSegment()
        speechSegmentList.append(nextSpeechSegment)
        nextSpeechSegment.segmentStartTime = float(turn["start_timestamp_millis"]) / 1000.0
        nextSpeechSegment.segmentEndTime = float(turn["end_timestamp_millis"]) / 1000.0
        nextSpeechSegment.segmentSpeaker = turn["speaker"]["speaker_label"]
        nextSpeechSegment.segmentText = turn["text"]
        nextSpeechSegment.segmentLanguage = turn["language"].lower()
        confidenceList = []
        nextSpeechSegment.segmentConfidence = confidenceList
        skipLeadingSpace = True

        # Process each word in this turn
        for word_ptr in turn["audio_item_indices"]:
            # Pick out our next data from a 'pronunciation'
            word = data["audio_items"][word_ptr]
            if "start_timestamp_millis" in word:
                # Write the word, and a leading space if this isn't the start of the segment
                if skipLeadingSpace:
                    skipLeadingSpace = False
                    wordToAdd = word["content"]
                else:
                    wordToAdd = " " + word["content"]

                # Add the word and confidence to this segment's list
                confidenceList.append({"text": wordToAdd,
                                        "confidence": 1.0,
                                        "start_time": float(word["start_timestamp_millis"]) / 1000.0,
                                        "end_time": float(word["end_timestamp_millis"] / 1000.0)})
            else:
                # Punctuation, needs to be added to the previous word
                last_word = nextSpeechSegment.segmentConfidence[-1]
                last_word["text"] = last_word["text"] + word["content"]

        # Extract guardrail information that breaced our limt
        if cli_args.guardrailCheck == 'on':
            for guardrail in data["audio"]["content_moderation"][index]["moderation_categories"]:
                if float(guardrail["confidence"]) >= float(cli_args.guardrailLimit):
                    nextSpeechSegment.segmentContentModeration.append(guardrail)
        index += 1

    # Return our full turn-by-turn speaker segment list with sentiment,
    # along with a flag to indicate the presence of call summary data
    speechSegmentList = merge_speaker_segments(speechSegmentList)
    return speechSegmentList


def generate_document():
    """
    Entrypoint for the command-line interface.
    """
    # Parameter extraction
    cli_parser = argparse.ArgumentParser(prog='bda-to-word',
                                         description='Turn an Amazon Bedrock Data Autimation Audio job output into an MS Word document')
    cli_parser.add_argument('--inputFile', metavar='filename', type=str, help='File containing BDA output JSON output', required=True)
    cli_parser.add_argument('--outputFile', metavar='filename', type=str, help='Output file to hold MS Word document')
    cli_parser.add_argument('--guardrailCheck', choices=['on', 'off'], default='off', help='Enables or disable reporting of guardrail breaches')
    cli_parser.add_argument('--guardrailLimit', type=float, default=0.2, help='Threshold limit for reporting guardrail breaches')
    cli_parser.add_argument('--sentiment', choices=['on', 'off'], default='off', help='Enables sentiment analysis on each conversational turn via Amazon Comprehend')
    cli_parser.add_argument('--customFile', metavar='filename', type=str, help='File containing standard BDA Audio custom blueprint output')
    cli_args = cli_parser.parse_args()

    # Load in the JSON file for processing
    json_filepath = Path(cli_args.inputFile)
    if json_filepath.is_file():
        json_data = json.load(open(json_filepath.absolute(), "r", encoding="utf-8"))
    else:
        print("FAIL: Specified JSON file '{0}' does not exist.".format(cli_args.inputFile))
        exit(-1)

    # Do something with the outputFile
    if cli_args.outputFile is None:
        cli_args.outputFile = cli_args.inputFile + ".docx"

    # Flag to indicate we have custom file output too
    custom_json_data = {}
    if cli_args.customFile is not None:
        custom_json_filepath = Path(cli_args.customFile)
        if custom_json_filepath.is_file():
            custom_json_data = json.load(open(custom_json_filepath.absolute(), "r", encoding="utf-8"))
        else:
            print("FAIL: Specified custom JSON file '{0}' does not exist.".format(cli_args.customFile))
            exit(-1)

    # May as well disable guardrail checking if it's off
    if cli_args.guardrailCheck == 'on':
        if "content_moderation" not in json_data['audio']:
            cli_args.guardrailCheck == 'off'

    # Generate the core transcript
    start = perf_counter()
    speech_segments = create_turn_by_turn_segments(json_data, cli_args)

    # Inject Comprehend-based sentiments into the segment list if required
    if cli_args.sentiment == 'on':
        generate_sentiment(speech_segments)

    # Write out our file and the performance statistics
    write(cli_args, speech_segments, custom_json_data)
    finish = perf_counter()
    duration = round(finish - start, 2)
    print(f"> Transcript {cli_args.outputFile} writen in {duration} seconds.")

# Main entrypoint
if __name__ == "__main__":
    generate_document()
